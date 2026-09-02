#!/usr/bin/env python3
"""Build editable, one-column Word manuscripts from the canonical Markdown files."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
RESULTS = ROOT / "results"

MAIN_TABLE_CAPTIONS = {
    1: "Data completeness and descriptive statistics.",
    2: "Correlation and collinearity screening.",
    3: "Heatwave definition and reference-period sensitivity.",
    4: "Temperature-trend estimates and sensitivity analyses.",
    5: "Poisson and negative-binomial model comparison.",
    6: "Primary negative-binomial count-model estimate.",
    7: "Adjusted antecedent meteorological associations and sensitivity analyses.",
    8: "Strictly chronological blocked-validation performance.",
}


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _clean_inline(text: str) -> str:
    """Remove Markdown-only decoration while retaining LaTeX source as editable text."""
    return text.replace("**", "").replace("`", "")


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(0, 0, 0)
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].paragraph_format.space_after = Pt(6)


def _blacken(paragraph) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_markdown(document: Document, path: Path, highlighted: bool = False) -> None:
    """Add simple Markdown as editable Word paragraphs; figures are appended separately."""
    skip_captions = False
    references = False
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line == "## Figure captions":
            skip_captions = True
            continue
        if skip_captions and line == "## References":
            skip_captions = False
            references = True
        if skip_captions:
            continue
        if not line:
            document.add_paragraph()
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(_clean_inline(line[2:]), style="Title")
        elif line.startswith("## "):
            paragraph = document.add_paragraph(_clean_inline(line[3:]), style="Heading 1")
        elif line.startswith("### "):
            paragraph = document.add_paragraph(_clean_inline(line[4:]), style="Heading 2")
        elif line.startswith("- "):
            paragraph = document.add_paragraph(_clean_inline(line[2:]), style="List Bullet")
        else:
            paragraph = document.add_paragraph(_clean_inline(line))
        _blacken(paragraph)
        substantive = (
            highlighted
            and paragraph.style.name in {"Normal", "List Bullet"}
            and not references
            and not line.startswith("[")
        )
        if substantive:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_dataframe(document: Document, frame: pd.DataFrame, caption: str) -> None:
    paragraph = document.add_paragraph(caption, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _blacken(paragraph)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.autofit = True
    for index, column in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        cell.text = str(column).replace("_", " ")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0, 0, 0)
    for values in frame.fillna("").itertuples(index=False, name=None):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if isinstance(value, float):
                value = f"{value:.5g}"
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7)
                    run.font.color.rgb = RGBColor(0, 0, 0)
    document.add_paragraph()


def _extract_figure_captions() -> dict[int, str]:
    captions: dict[int, str] = {}
    for line in (MANUSCRIPT / "original_article_clean.md").read_text(encoding="utf-8").splitlines():
        match = re.match(r"\*\*Figure (\d+)\. (.*?)\*\*\s*(.*)", line)
        if match:
            number, title, detail = match.groups()
            captions[int(number)] = f"Figure {number}. {title}. {detail}"
    return captions


def _add_main_tables_and_figures(document: Document) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Tables", level=1)
    for number, caption in MAIN_TABLE_CAPTIONS.items():
        table_dir = RESULTS / "tables" / "main"
        matches = sorted(table_dir.glob(f"main_table{number:02d}_*.csv"))
        if len(matches) != 1:
            raise RuntimeError(f"Expected one main Table {number}, found {matches}")
        _add_dataframe(document, pd.read_csv(matches[0]), f"Table {number}. {caption}")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Figures", level=1)
    captions = _extract_figure_captions()
    for number in range(1, 8):
        matches = sorted((RESULTS / "figures" / "main").glob(f"figure{number:02d}_*.png"))
        if len(matches) != 1:
            raise RuntimeError(f"Expected one main Figure {number}, found {matches}")
        document.add_picture(str(matches[0]), width=Inches(6.35))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph(captions[number], style="Caption")
        _blacken(paragraph)
        document.add_paragraph()


def build_article(output_name: str, highlighted: bool) -> None:
    document = Document()
    _style_document(document)
    _add_markdown(document, MANUSCRIPT / "original_article_clean.md", highlighted=highlighted)
    _add_main_tables_and_figures(document)
    document.core_properties.title = "Long-Term Warming and Definition-Dependent Heatwaves in Dhaka"
    document.core_properties.subject = "Editable original-article manuscript"
    document.save(MANUSCRIPT / output_name)


def build_supplement() -> None:
    document = Document()
    _style_document(document)
    _add_markdown(document, MANUSCRIPT / "supplementary_material.md")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Supplementary tables", level=1)
    supplement_tables = [
        ("supplement_tableS01_count_influence_sensitivity.csv", "Supplementary Table S1. Leave-one-influential-year-out NB2 count sensitivity."),
        ("supplement_tableS02_forecast_validation.csv", "Supplementary Table S2. Forecast metrics by rolling origin and model."),
    ]
    for name, caption in supplement_tables:
        _add_dataframe(document, pd.read_csv(RESULTS / "tables" / "supplement" / name), caption)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Supplementary figures", level=1)
    supplement_figures = [
        ("figureS01_selected_pairplots.png", "Supplementary Figure S1. Selected meteorological pairplots. Exploratory hot-season distributions colored by primary persistent-day status."),
        ("figureS02_forecast_validation.png", "Supplementary Figure S2. Rolling-origin forecast validation. Accuracy, observed-versus-predicted values, interval coverage, and interval width."),
    ]
    for name, caption in supplement_figures:
        document.add_picture(str(RESULTS / "figures" / "supplement" / name), width=Inches(6.35))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph(caption, style="Caption")
        _blacken(paragraph)
    document.save(MANUSCRIPT / "supplementary_material.docx")


def main() -> None:
    build_article("original_article_clean.docx", highlighted=False)
    build_article("original_article_updates_highlighted_yellow.docx", highlighted=True)
    build_supplement()
    print("Built clean, yellow-highlighted, and supplementary editable Word documents.")


if __name__ == "__main__":
    main()
